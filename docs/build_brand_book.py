"""
Generates Won Vision Brand Book.docx — a designer-facing brand
document with the mark embedded, real colour swatches as filled
table cells, and Sora type specimens.

Run with:
    uv run --with python-docx python docs/build_brand_book.py
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ---------------------------------------------------------------- paths
ROOT = Path(__file__).resolve().parents[1]
LOGO_PNG = Path("/Volumes/WON SOL/WON VISION DESIGN/Won Vision Symbol.png")
OUT_REPO = ROOT / "docs" / "Won Vision Brand Book.docx"
OUT_DRIVE = Path("/Volumes/WON SOL/WON VISION DESIGN/Won Vision Brand Book.docx")

# ---------------------------------------------------------------- helpers
def shade_cell(cell, hex_no_hash: str):
    """Fill a table cell with a flat colour."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_no_hash)
    tc_pr.append(shd)


def remove_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tbl_pr.append(borders)


def set_run_font(run, name="Sora", size_pt=11, bold=False, color_hex=None,
                 letter_spacing_pt=None, all_caps=False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex.lstrip("#"))
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), name)
    if letter_spacing_pt is not None:
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:val"), str(int(letter_spacing_pt * 20)))
        rpr.append(spacing)
    if all_caps:
        caps = OxmlElement("w:caps")
        caps.set(qn("w:val"), "true")
        rpr.append(caps)


def add_para(doc, text="", *, size=11, bold=False, color="#000000",
             space_before=0, space_after=4, align=None, all_caps=False,
             letter_spacing_pt=None, line_spacing=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p_fmt = p.paragraph_format
    p_fmt.space_before = Pt(space_before)
    p_fmt.space_after = Pt(space_after)
    if line_spacing is not None:
        p_fmt.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        set_run_font(run, size_pt=size, bold=bold, color_hex=color,
                     letter_spacing_pt=letter_spacing_pt, all_caps=all_caps)
    return p


def add_eyebrow(doc, text):
    add_para(doc, text, size=8, bold=True, color="#737373",
             letter_spacing_pt=2.2, all_caps=True, space_before=18,
             space_after=8)


def add_section_title(doc, label, title):
    add_eyebrow(doc, label)
    add_para(doc, title, size=32, bold=True, all_caps=True,
             letter_spacing_pt=-0.7, line_spacing=1.0,
             space_before=0, space_after=10)


def hr(doc, color="#000000"):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), color.lstrip("#"))
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    p.paragraph_format.space_after = Pt(8)


def page_break(doc):
    doc.add_page_break()


# ---------------------------------------------------------------- doc
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)

# default body font
style = doc.styles["Normal"]
style.font.name = "Sora"
style.font.size = Pt(11)

# ============================================================ COVER
if LOGO_PNG.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(140)
    p.paragraph_format.space_after = Pt(40)
    p.add_run().add_picture(str(LOGO_PNG), width=Inches(3.4))

add_para(doc, "Won Vision", size=56, bold=True, all_caps=True,
         letter_spacing_pt=-1.4, align=WD_ALIGN_PARAGRAPH.CENTER,
         line_spacing=0.95, space_after=14)
add_para(doc, "Brand System v2", size=14, bold=False, color="#404040",
         all_caps=True, letter_spacing_pt=2.0,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=80)
add_para(doc, "Real estate photography studio. Melbourne.",
         size=11, color="#737373", align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Won Vision Pty Ltd · May 2026",
         size=8, color="#737373", letter_spacing_pt=2.2, all_caps=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8)

page_break(doc)

# ============================================================ INTRO
add_section_title(doc, "01 — The idea", "The brand")
hr(doc)
add_para(
    doc,
    "Won Vision is a Melbourne real-estate photography studio that lives "
    "on craft, restraint, and seeing what other photographers miss.",
    size=11, line_spacing=1.5, space_after=10,
)
add_para(
    doc,
    "The brand is monochrome and quiet. Black and white type, sharp "
    "corners, uncomplicated layouts. The work — the photographs — are "
    "the only place colour ever appears, and they're shown unedited "
    "and full-saturation.",
    size=11, line_spacing=1.5, space_after=10,
)
add_para(
    doc,
    "If a piece of design starts feeling busy, decorative, or "
    "“modern”, it isn't Won Vision.",
    size=11, line_spacing=1.5,
)

page_break(doc)

# ============================================================ THE MARK
add_section_title(doc, "02 — The mark", "The W + V monogram")
hr(doc)

if LOGO_PNG.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(20)
    p.add_run().add_picture(str(LOGO_PNG), width=Inches(2.6))

add_para(doc, "Anatomy", size=12, bold=True, space_before=12, space_after=6)
add_para(
    doc,
    "A W and a V interlocked through a shared centre stroke, with a "
    "circular dot top-right. The triangular negative spaces inside the "
    "W and V read as aperture cuts — the mark's tie to photography. "
    "The dot is the only round element in the entire system.",
    size=11, line_spacing=1.5, space_after=14,
)

add_para(doc, "Sizing reference", size=12, bold=True, space_after=8)

sizes = [
    ("Hero / about page", "80 — 120 px"),
    ("Website navigation", "32 — 40 px"),
    ("Footer", "24 — 32 px"),
    ("Favicon / app icon", "32 px (mark only)"),
    ("Watermark on photos", "24 — 48 px"),
    ("Business card", "18 — 24 mm"),
]
t = doc.add_table(rows=len(sizes), cols=2)
remove_table_borders(t)
t.columns[0].width = Cm(7.5)
t.columns[1].width = Cm(6.5)
for i, (where, size_) in enumerate(sizes):
    c1, c2 = t.rows[i].cells
    c1.text = where
    c2.text = size_
    set_run_font(c1.paragraphs[0].runs[0], size_pt=10, color_hex="#404040")
    set_run_font(c2.paragraphs[0].runs[0], size_pt=10, bold=True)
    if i % 2 == 0:
        shade_cell(c1, "F5F5F5")
        shade_cell(c2, "F5F5F5")

add_para(doc, "Clearspace", size=12, bold=True, space_before=18, space_after=6)
add_para(
    doc,
    "Always leave at least 25% of the mark's height as empty space on "
    "every side. The dot needs room to breathe — the negative space "
    "around it is part of the design.",
    size=11, line_spacing=1.5, space_after=12,
)

add_para(doc, "Wordmark lockup", size=12, bold=True, space_before=8, space_after=6)
add_para(
    doc,
    "When the mark sits next to the words “Won Vision”:",
    size=11, line_spacing=1.5, space_after=4,
)
for line in [
    "Mark on the left, words on the right.",
    "“Won Vision” is sentence case — never WON VISION.",
    "Set in Sora, weight 500.",
    "Mark height ≈ 1.4× the height of the letters.",
    "Gap between mark and words ≈ 0.4× mark height.",
]:
    add_para(doc, "  •  " + line, size=11, line_spacing=1.5, space_after=2)

add_para(doc, "Mark — never", size=12, bold=True, space_before=18, space_after=6)
for line in [
    "Recolour outside black, white, or palette greys.",
    "Add stroke, outline, drop shadow, glow, or gradient.",
    "Crop the dot.",
    "Separate the W from the V.",
    "Rotate or skew the mark.",
    "Render “Won Vision” in all caps.",
]:
    add_para(doc, "  ✕  " + line, size=11, color="#404040",
             line_spacing=1.5, space_after=2)

page_break(doc)

# ============================================================ COLOUR
add_section_title(doc, "03 — Colour", "Pure black and white")
hr(doc)

add_para(
    doc,
    "Pure black and white. No accent colours. Photographs run in "
    "full colour, unedited — they are the only place hue exists in "
    "the system.",
    size=11, line_spacing=1.5, space_after=14,
)

palette = [
    # (label, hex, role, dark text)
    ("Background",        "#FFFFFF", "Default page background", True),
    ("Foreground",        "#000000", "Headings, body, mark",     False),
    ("Inverse background","#000000", "Hero, drawer, lightbox",   False),
    ("Inverse foreground","#FFFFFF", "Text on dark sections",    True),
    ("Text — secondary",  "#404040", "Sub-headings, body emphasis", False),
    ("Text — muted",      "#737373", "Captions, eyebrows, meta", False),
    ("Border light",      "#E5E5E5", "Dividers on white",        True),
    ("Border strong",     "#999999", "High-contrast dividers",   False),
    ("Border inverse",    "#2A2A2A", "Dividers on black",        False),
]

t = doc.add_table(rows=len(palette), cols=4)
remove_table_borders(t)
t.columns[0].width = Cm(2.5)
t.columns[1].width = Cm(4.5)
t.columns[2].width = Cm(3.0)
t.columns[3].width = Cm(5.5)

for i, (label, hex_, role, dark_label) in enumerate(palette):
    c_swatch, c_label, c_hex, c_role = t.rows[i].cells
    # set row height
    tr = t.rows[i]._tr
    tr_pr = tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), "900")  # twentieths of a point ≈ 1.6cm
    tr_pr.append(tr_height)

    # swatch cell — fill with the actual colour
    shade_cell(c_swatch, hex_.lstrip("#"))
    c_swatch.text = ""

    # label
    c_label.text = label
    set_run_font(c_label.paragraphs[0].runs[0], size_pt=11, bold=True)
    c_label.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # hex
    c_hex.text = hex_
    set_run_font(c_hex.paragraphs[0].runs[0], size_pt=10, color_hex="#404040")
    c_hex.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # role
    c_role.text = role
    set_run_font(c_role.paragraphs[0].runs[0], size_pt=10, color_hex="#737373")
    c_role.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

add_para(doc, "Colour rules", size=12, bold=True, space_before=22, space_after=6)
for line in [
    "Default mode is white background, black foreground. Inverse "
    "(black on white) is reserved for high-impact moments — the hero, "
    "drawer, and lightbox.",
    "Photographs always render in full colour, unedited. Don't apply "
    "a black-and-white filter, tint, or wash.",
    "No blue links, green CTAs, or red errors. Hierarchy uses weight, "
    "size, and grey value — never hue.",
    "Don't pair grey on grey for primary content. Always pure-on-pure "
    "for headings and body.",
]:
    add_para(doc, "  •  " + line, size=11, line_spacing=1.5, space_after=4)

page_break(doc)

# ============================================================ TYPOGRAPHY
add_section_title(doc, "04 — Typography", "Sora")
hr(doc)

add_para(
    doc,
    "The brand uses one typeface — Sora — at two weights. There is no "
    "italic. Emphasis comes from weight, size, and uppercase eyebrows.",
    size=11, line_spacing=1.5, space_after=14,
)

# weight specimen — display
add_eyebrow(doc, "Sora 500 · Display & headings")
add_para(doc, "Won Vision", size=48, bold=True,
         all_caps=True, letter_spacing_pt=-1.2, line_spacing=0.95,
         space_after=14)

# weight specimen — body
add_eyebrow(doc, "Sora 400 · Body & captions")
add_para(
    doc,
    "Won Vision is a Melbourne real-estate photography studio. "
    "Photography, listing video, CASA-licensed drone, floor plans, "
    "virtual staging, and headshots — flexible packages and add-ons "
    "built to sell premium property faster.",
    size=12, line_spacing=1.55, space_after=18,
)

# size table
add_para(doc, "Type scale", size=12, bold=True, space_before=8, space_after=8)

scale = [
    ("Display", "56 — 140 px", "All caps", 28, True, True),
    ("H1",      "40 — 80 px",  "All caps", 22, True, True),
    ("H2",      "28 — 48 px",  "Sentence case", 18, True, False),
    ("H3",      "20 — 28 px",  "Sentence case", 14, True, False),
    ("Body lg", "18 — 22 px",  "Sentence case", 12, False, False),
    ("Body",    "16 px",        "Sentence case", 11, False, False),
    ("Caption", "14 px",        "Sentence case", 10, False, False),
    ("Eyebrow", "11 px",        "UPPERCASE +0.2em", 9,  True, True),
]

t = doc.add_table(rows=len(scale)+1, cols=4)
remove_table_borders(t)
t.columns[0].width = Cm(3.0)
t.columns[1].width = Cm(3.5)
t.columns[2].width = Cm(4.5)
t.columns[3].width = Cm(4.5)

# header row
hdr = ["Level", "Size", "Treatment", "Specimen"]
for j, h in enumerate(hdr):
    cell = t.rows[0].cells[j]
    cell.text = h
    set_run_font(cell.paragraphs[0].runs[0], size_pt=8, bold=True,
                 color_hex="#737373", letter_spacing_pt=2.2,
                 all_caps=True)
    shade_cell(cell, "F5F5F5")

for i, (lv, sz, tr, sp_pt, bold, caps) in enumerate(scale, start=1):
    cells = t.rows[i].cells
    cells[0].text = lv
    set_run_font(cells[0].paragraphs[0].runs[0], size_pt=10, bold=True)
    cells[1].text = sz
    set_run_font(cells[1].paragraphs[0].runs[0], size_pt=10,
                 color_hex="#404040")
    cells[2].text = tr
    set_run_font(cells[2].paragraphs[0].runs[0], size_pt=10,
                 color_hex="#737373")
    # specimen
    cells[3].text = lv
    set_run_font(cells[3].paragraphs[0].runs[0], size_pt=sp_pt, bold=bold,
                 all_caps=caps,
                 letter_spacing_pt=-0.6 if sp_pt >= 18 else None)

add_para(doc, "Type rules", size=12, bold=True, space_before=22, space_after=6)
for line in [
    "Display headings and H1 are ALL CAPS. They carry the brand's "
    "energy and match the geometry of the mark.",
    "H2 and H3 are sentence case.",
    "Body copy is never uppercase. Only eyebrows and small labels are.",
    "The wordmark “Won Vision” is sentence case — the one "
    "exception to the display all-caps rule.",
    "Letter-spacing: tight on display (-0.03em), neutral on body, "
    "wide on eyebrows (+0.2em).",
    "Line-height: tight on display (0.95), comfortable on body (1.5).",
]:
    add_para(doc, "  •  " + line, size=11, line_spacing=1.5, space_after=4)

# ----- Getting Sora ---------------------------------------------------
add_para(doc, "Getting Sora", size=12, bold=True,
         space_before=22, space_after=6)
add_para(
    doc,
    "Sora is free and open-source — released under the SIL Open Font "
    "Licence by Google Fonts. Use it on every Won Vision surface "
    "(decks, documents, social, video titles, signage). No purchase, "
    "no licence key.",
    size=11, line_spacing=1.5, space_after=10,
)

add_para(doc, "Download (one-time, install on your computer)",
         size=11, bold=True, space_before=8, space_after=4)
for line in [
    "Open  fonts.google.com/specimen/Sora",
    "Click  “Get font”  →  “Download all”.",
    "Unzip the file. You'll see “static/” and the variable file "
    "“Sora-VariableFont_wght.ttf”.",
    "macOS: double-click the .ttf, then click “Install Font” in Font "
    "Book. The variable file alone covers every weight.",
    "Windows: right-click the .ttf and choose “Install for all users”.",
    "Restart Word, Pages, Keynote, Figma — Sora will appear in their "
    "font menus.",
]:
    add_para(doc, "  •  " + line, size=11, line_spacing=1.5, space_after=3)

add_para(doc, "Use on the web",
         size=11, bold=True, space_before=12, space_after=4)
add_para(
    doc,
    "The website already loads Sora through Next.js. For any other "
    "web project, paste this into the <head> of the page:",
    size=11, line_spacing=1.5, space_after=6,
)
code = (
    "<link rel=\"preconnect\" href=\"https://fonts.googleapis.com\">\n"
    "<link rel=\"preconnect\" href=\"https://fonts.gstatic.com\" crossorigin>\n"
    "<link href=\"https://fonts.googleapis.com/css2?family=Sora:wght@400;500&display=swap\" rel=\"stylesheet\">"
)
p = doc.add_paragraph(code)
p.paragraph_format.left_indent = Cm(0.6)
p.paragraph_format.space_after = Pt(8)
for r in p.runs:
    set_run_font(r, name="Menlo", size_pt=9, color_hex="#404040")

add_para(
    doc,
    "Then reference it in CSS as:  font-family: \"Sora\", sans-serif;",
    size=11, line_spacing=1.5, space_after=8,
)

add_para(doc, "Use in Microsoft Office / Google Workspace",
         size=11, bold=True, space_before=12, space_after=4)
for line in [
    "Word, PowerPoint, Excel — once Sora is installed on the system, "
    "it appears in the font menu like any other font.",
    "Google Docs / Slides — open the font menu, choose “More fonts”, "
    "type “Sora”, tick it. Available in every Workspace document "
    "from then on.",
    "Canva — search “Sora” in the font picker. Free, no upload "
    "needed.",
]:
    add_para(doc, "  •  " + line, size=11, line_spacing=1.5, space_after=3)

add_para(
    doc,
    "Substitute only if Sora truly isn't available — in that case "
    "fall back to the system sans-serif (San Francisco on macOS, "
    "Segoe UI on Windows). Never substitute another Google font.",
    size=10, color="#737373", line_spacing=1.5, space_before=10,
    space_after=4,
)

page_break(doc)

# ============================================================ PHOTO
add_section_title(doc, "05 — Photography", "The work is the colour")
hr(doc)

for line in [
    "Naturally coloured. No preset, no over-direction, no "
    "orange-and-teal.",
    "Twilight on request only — not a default look.",
    "Wide architectural and tight detail in the same pass.",
    "Aerial integrates with ground; don't deliver as a separate "
    "aesthetic.",
    "Floor plans render in the brand's colours (black ink on white).",
]:
    add_para(doc, "  •  " + line, size=11, line_spacing=1.5, space_after=4)

add_para(doc, "Presenting photographs", size=12, bold=True,
         space_before=18, space_after=6)
for line in [
    "Square corners only — never round photo edges.",
    "No filters, tints, or drop shadows in marketing surfaces.",
    "Full bleed where possible. Don't crop a great photo into a small "
    "card.",
    "Captions sit below or over the bottom of the image, never inside "
    "the frame.",
    "A dark gradient veil at the bottom is acceptable when text overlays "
    "the photo — neutral black, never tinted.",
]:
    add_para(doc, "  •  " + line, size=11, line_spacing=1.5, space_after=4)

page_break(doc)

# ============================================================ VOICE
add_section_title(doc, "06 — Voice", "How Won Vision sounds")
hr(doc)

add_para(doc, "Confident, never loud.", size=22, bold=True, all_caps=False,
         letter_spacing_pt=-0.5, line_spacing=1.2, space_after=4)
add_para(doc, "Precise, never vague.", size=22, bold=True,
         letter_spacing_pt=-0.5, line_spacing=1.2, space_after=4)
add_para(doc, "Clear, never clever.", size=22, bold=True,
         letter_spacing_pt=-0.5, line_spacing=1.2, space_after=22)

add_para(doc, "Avoid", size=12, bold=True, space_after=6)
for line in [
    "Real-estate jargon (“nestled”, “boasts”, "
    "“stunning”).",
    "Marketing superlatives (“the best”, “the only”, "
    "“the ultimate”).",
    "Emoji in copy or UI.",
    "Hashtags in body copy. Save them for social posts, used sparingly.",
]:
    add_para(doc, "  •  " + line, size=11, line_spacing=1.5, space_after=4)

page_break(doc)

# ============================================================ DON'TS
add_section_title(doc, "07 — Brand-wide don’ts", "Hard rules")
hr(doc)

donts = [
    "No accent colours — no blue links, no green CTAs, no red errors.",
    "No other typefaces — Sora only.",
    "No effects on the logo or text — no gradients, shadows, glows, or blurs.",
    "No greyscale or filtered photos. The work is in colour.",
    "No rounded photo corners. Square only.",
    "No border-radius greater than 4px anywhere.",
    "No emoji in any user interface.",
    "No italic body type — emphasis comes from weight.",
    "No “WON VISION” in caps — the wordmark is always sentence case.",
    "No centred-template layouts.",
]
for line in donts:
    add_para(doc, "  ✕  " + line, size=12, color="#000000",
             line_spacing=1.5, space_after=6)

page_break(doc)

# ============================================================ QUICK REF
add_section_title(doc, "08 — Quick reference", "At a glance")
hr(doc)

ref = [
    ("Mark", "W + V monogram + dot"),
    ("Foreground", "#000000"),
    ("Background", "#FFFFFF"),
    ("Inverse", "#FFFFFF on #000000"),
    ("Type family", "Sora"),
    ("Heading & logo weight", "500"),
    ("Body weight", "400"),
    ("Display case", "UPPERCASE"),
    ("Heading case", "Sentence case"),
    ("Wordmark case", "Sentence case (Won Vision)"),
    ("Border radius", "≤ 4px"),
    ("Photo corners", "Square"),
    ("Photo treatment", "Full colour, unedited"),
    ("Loading animation", "≈ 3 seconds, draws then holds"),
    ("Mobile breakpoint", "900 px"),
]
t = doc.add_table(rows=len(ref), cols=2)
remove_table_borders(t)
t.columns[0].width = Cm(7.5)
t.columns[1].width = Cm(8.0)
for i, (k, v) in enumerate(ref):
    c1, c2 = t.rows[i].cells
    c1.text = k
    c2.text = v
    set_run_font(c1.paragraphs[0].runs[0], size_pt=10, bold=True,
                 color_hex="#404040", letter_spacing_pt=1.0,
                 all_caps=True)
    set_run_font(c2.paragraphs[0].runs[0], size_pt=11)
    if i % 2 == 0:
        shade_cell(c1, "FAFAFA")
        shade_cell(c2, "FAFAFA")

add_para(doc, "Won Vision Pty Ltd · Brand System v2 · May 2026",
         size=8, color="#737373", letter_spacing_pt=2.2, all_caps=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60)

# ---------------------------------------------------------------- save
OUT_REPO.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT_REPO)
print(f"Saved: {OUT_REPO}")

if OUT_DRIVE.parent.exists():
    doc.save(OUT_DRIVE)
    print(f"Saved: {OUT_DRIVE}")
else:
    print(f"Drive not mounted — skipped {OUT_DRIVE}")
